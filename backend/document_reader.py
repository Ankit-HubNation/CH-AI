from pathlib import Path
from typing import Any

from docx import Document


SUPPORTED_EXTENSIONS = {".docx"}


def document_roots(project_dir: Path) -> list[Path]:
    return [project_dir / "documents", project_dir / "uploads"]


def safe_document_path(project_dir: Path, filename: str) -> Path:
    requested = Path(filename).name
    for root in document_roots(project_dir):
        candidate = root / requested
        if candidate.is_file() and candidate.suffix.lower() in SUPPORTED_EXTENSIONS:
            return candidate
    raise FileNotFoundError(requested)


def extract_docx(path: Path) -> str:
    document = Document(path)
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)


def list_docx_files(project_dir: Path) -> list[dict[str, Any]]:
    files: list[dict[str, Any]] = []
    seen: set[str] = set()
    for root in document_roots(project_dir):
        root.mkdir(parents=True, exist_ok=True)
        for path in sorted(root.glob("*.docx")):
            if path.name in seen:
                continue
            seen.add(path.name)
            text = extract_docx(path)
            files.append(
                {
                    "filename": path.name,
                    "source": root.name,
                    "size_bytes": path.stat().st_size,
                    "paragraphs": len(text.split("\n\n")) if text else 0,
                    "content": text,
                }
            )
    return files
